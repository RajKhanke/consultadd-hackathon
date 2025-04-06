import pandas as pd
import PyPDF2
import os
from typing import Dict, List, Tuple, Optional
import json
import os
import sys
import asyncio
from dataclasses import dataclass
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
from pydantic import BaseModel, Field
import docx

from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from langchain.retrievers.document_compressors import CohereRerank
from langchain.retrievers import ContextualCompressionRetriever
from langchain.schema import Document

from pydantic_ai import Agent, RunContext
from pydantic_ai.models.groq import GroqModel
from pydantic_ai.providers.groq import GroqProvider

# Load environment variables
load_dotenv()

# Constants
FAISS_INDEX_FOLDER = os.path.join(os.path.expanduser("~"), "OneDrive", "Documents", "consultadd", "rag_data", "faiss_indexes")
MAX_CHUNKS_SUMMARY = 25  # Maximum chunks for summarization
MAX_CHUNKS_ELIGIBILITY = 50  # Maximum chunks for eligibility checks
RERANKING_ENABLED = False  # Set to True if you have a Cohere API key

# Check for required environment variables
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY not found in environment variables")

# Check for optional Cohere API key for reranking
COHERE_API_KEY = os.environ.get("COHERE_API_KEY")
if COHERE_API_KEY:
    RERANKING_ENABLED = True
    os.environ["COHERE_API_KEY"] = COHERE_API_KEY
else:
    print("Warning: COHERE_API_KEY not found. Reranking will be disabled.")

# Pydantic models for structured outputs
class DocumentSummary(BaseModel):
    """Document summary model"""
    context: str = Field(..., description="Overall context and purpose of the document")
    scope_of_work: str = Field(..., description="Description of the work or services being requested")
    key_points: List[str] = Field(..., description="Key points about the document")
    document_type: str = Field(..., description="Type of the document (RFP, proposal, contract, etc.)")

class EligibilityCriteria(BaseModel):
    """Eligibility criteria model"""
    mandatory_requirements: List[str] = Field(..., description="List of all mandatory requirements")
    qualifications: List[str] = Field(..., description="Required qualifications")
    certifications: List[str] = Field(..., description="Required certifications")
    experience: List[str] = Field(..., description="Required experience")
    missing_criteria: List[str] = Field(..., description="Potentially missing or unclear criteria")
    recommendation: str = Field(..., description="Overall assessment of the eligibility requirements")

# New Pydantic models for compliance check
class ComplianceCheckResult(BaseModel):
    """Compliance check result model"""
    overall_score: float = Field(..., description="Overall compliance score (0-100)")
    eligibility_score: float = Field(..., description="Legal eligibility score (0-100)")
    domain_match_score: float = Field(..., description="Domain alignment score (0-100)")
    certifications_score: float = Field(..., description="Certifications compliance score (0-100)")
    experience_score: float = Field(..., description="Experience requirements score (0-100)")
    deal_breakers: List[str] = Field(..., description="Potential deal-breakers identified")
    recommendations: List[str] = Field(..., description="Recommendations for compliance")

@dataclass
class ComplianceDeps:
    """Dependencies for the compliance check agent"""
    collection_name: str
    company_data: Dict[str, Any] 

COMPANY_DATA = {}

# Function to load and process company data
def load_company_data(file_path: str) -> Dict[str, Any]:
    """
    Load and process company data from the provided file.
    
    Args:
        file_path: Path to the company data file.
        
    Returns:
        Dictionary with company data.
    """
    # Determine file type and extract text
    ext = file_path.lower().rsplit(".", 1)[-1]
    raw_text = ""
    
    try:
        if ext == "pdf":
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        raw_text += page_text + "\n"
        elif ext == "docx":
            doc = docx.Document(file_path)
            raw_text = "\n".join(para.text for para in doc.paragraphs)
        elif ext == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()
        elif ext == "json":
            with open(file_path, "r", encoding="utf-8") as f:
                return json.load(f)
        else:
            raise ValueError(f"Unsupported file format: .{ext}")
    except Exception as e:
        print(f"Error loading company data: {str(e)}")
        return {}
    
    # Parse the text into key-value pairs
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    company_data = {}
    
    # Simple parsing strategy: look for key: value patterns
    for line in lines:
        if ":" in line:
            parts = line.split(":", 1)
            key = parts[0].strip()
            value = parts[1].strip() if len(parts) > 1 else ""
            company_data[key] = value
    
    # Ensure key fields are present with defaults if not found
    required_fields = [
        "Company Name", "Industry", "Certifications", "Years in Business",
        "State Registrations", "Key Specializations", "Past Performance", 
        "Team Size", "Largest Contract Value"
    ]
    
    for field in required_fields:
        if field not in company_data:
            company_data[field] = "Not specified"
    
    return company_data

# Document retrieval functions
async def retrieve_document_chunks(collection_name: str, is_summary: bool = True) -> str:
    """
    Retrieve document content based on task type.
    For summary: Uses first N chunks in order.
    """
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )
    index_path = os.path.join(FAISS_INDEX_FOLDER, collection_name)
    if not os.path.exists(index_path):
        return f"Error: FAISS index not found for collection '{collection_name}'."
    try:
        vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        docs = []
        for doc_id in vectorstore.docstore._dict:
            doc = vectorstore.docstore._dict[doc_id]
            if hasattr(doc, "page_content") and hasattr(doc, "metadata"):
                docs.append(doc)
        if is_summary:
            docs.sort(key=lambda x: x.metadata.get("chunk_id", 0))
            selected_docs = docs[:MAX_CHUNKS_SUMMARY]
        else:
            selected_docs = docs

        results = []
        for i, doc in enumerate(selected_docs):
            chunk_id = doc.metadata.get("chunk_id", i)
            source = doc.metadata.get("source", collection_name)
            results.append(f"## Document Section {i+1}\n**Source:** {source} | **Chunk:** {chunk_id}\n\n{doc.page_content}\n")
        return "\n\n".join(results)
    except Exception as e:
        import traceback
        print(f"Error retrieving document chunks: {str(e)}")
        print(traceback.format_exc())
        return f"Error retrieving document chunks: {str(e)}"

async def semantic_search_retrieval(collection_name: str, query: str, top_k: int = 10) -> str:
    """
    Perform semantic search on document collection and retrieve most relevant chunks.
    """
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )
    index_path = os.path.join(FAISS_INDEX_FOLDER, collection_name)
    if not os.path.exists(index_path):
        return f"Error: FAISS index not found for collection '{collection_name}'."
    try:
        vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        retriever_vectordb = vectorstore.as_retriever(search_kwargs={"k": top_k})
        all_docs = []
        for doc_id in vectorstore.docstore._dict:
            doc = vectorstore.docstore._dict[doc_id]
            if hasattr(doc, "page_content") and hasattr(doc, "metadata"):
                all_docs.append(Document(page_content=doc.page_content, metadata=doc.metadata))
        keyword_retriever = BM25Retriever.from_documents(all_docs)
        keyword_retriever.k = top_k
        ensemble_retriever = EnsembleRetriever(
            retrievers=[retriever_vectordb, keyword_retriever],
            weights=[0.7, 0.3]
        )
        if RERANKING_ENABLED:
            compressor = CohereRerank()
            final_retriever = ContextualCompressionRetriever(
                base_compressor=compressor,
                base_retriever=ensemble_retriever
            )
        else:
            final_retriever = ensemble_retriever
        docs = final_retriever.get_relevant_documents(query)
        results = []
        for i, doc in enumerate(docs):
            chunk_id = doc.metadata.get("chunk_id", i)
            source = doc.metadata.get("source", collection_name)
            results.append(f"## Document Section {i+1}\n**Source:** {source} | **Chunk:** {chunk_id}\n\n{doc.page_content}\n")
        if not results:
            return "No relevant document chunks found."
        return "\n\n".join(results)
    except Exception as e:
        import traceback
        print(f"Error in semantic search: {str(e)}")
        print(traceback.format_exc())
        return f"Error performing semantic search: {str(e)}"

# Dependency models for each agent
@dataclass
class SummaryDeps:
    collection_name: str

@dataclass
class EligibilityDeps:
    collection_name: str

@dataclass
class SubmissionChecklistDeps:
    collection_name: str

@dataclass
class ContractRiskDeps:
    collection_name: str

# Agent 1: Document Summary
def create_summary_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    # Added instruction to avoid any tool call formatting in the final answer.
    agent = Agent(
        model,
        system_prompt=(
              "You are a specialized document summarization assistant. Your task is to generate a comprehensive yet concise "
            "summary of a document. Follow these steps:\n"
            "1. Analyze the document content provided to understand its purpose and structure.\n"
            "2. Extract the core context, scope of work, and key points.\n"
            "3. Create a structured summary that includes:\n"
            "   - Overall context and purpose of the document\n"
            "   - Description of the work or services being requested\n"
            "   - Key points about requirements, deliverables, and important details\n"
            "   - Type of the document (RFP, contract, proposal, etc.)\n"
            "4. Present your summary in a clear, structured format.\n\n"
            "Your summary should be factual, objective, and based solely on the provided document content."
        )
    )
    @agent.tool
    async def get_document_content(context: RunContext[SummaryDeps]) -> str:
        collection_name = context.deps.collection_name
        return await retrieve_document_chunks(collection_name, is_summary=True)
    return agent

# Agent 2: Eligibility Criteria Extraction
def create_eligibility_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    agent = Agent(
        model,
        system_prompt=(
              "You are a specialized eligibility criteria extraction assistant. Your task is to analyze documents and identify "
            "all mandatory eligibility requirements. Follow these steps:\n"
            "1. Carefully review the document content provided.\n"
            "2. Extract and list all mandatory requirements, focusing on:\n"
            "   - Required qualifications and credentials\n"
            "   - Required certifications or licenses\n"
            "   - Minimum experience requirements\n"
            "   - Technical capabilities or resources needed\n"
            "3. Flag any missing or unclear eligibility criteria that should be specified but aren't.\n"
            "4. Provide a structured output with all mandatory requirements and flagged missing elements.\n"
            "5. Include a brief assessment of the overall eligibility requirements (strict, standard, minimal, etc.).\n\n"
            "Focus specifically on requirements using language like 'must have', 'required', 'shall', or 'mandatory'."
        )
    )
    @agent.tool
    async def search_document_content(context: RunContext[EligibilityDeps], search_query: str) -> str:
        collection_name = context.deps.collection_name
        if not search_query or search_query.strip() == "":
            search_query = "mandatory requirements qualifications certifications experience technical capabilities eligibility"
        return await semantic_search_retrieval(collection_name, search_query, top_k=10)
    return agent

# Agent 3: Submission Checklist Generation
def create_submission_checklist_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    agent = Agent(
        model,
        system_prompt=(
            "You are a specialized assistant for generating a submission checklist. Your task is to extract and structure the RFP submission requirements from a document. "
            "Follow these steps:\n"
            "1. Retrieve document sections relevant to submission requirements, including document format details (e.g., page limit, font type/size, line spacing, TOC requirements) and any specific attachments or forms using the provided tool if needed.\n"
            "2. Organize the extracted information into a structured checklist.\n"
            "3. Present your output under the top-level heading **Generating a Submission Checklist** with clear markdown headings and bullet points.\n"
            "IMPORTANT: Do not include any function call syntax or tool formatting in your final response."
        )
    )
    @agent.tool
    async def search_document_content(context: RunContext[SubmissionChecklistDeps], search_query: str) -> str:
        collection_name = context.deps.collection_name
        if not search_query or search_query.strip() == "":
            search_query = "submission checklist document format page limit font type line spacing attachments forms TOC requirements"
        return await semantic_search_retrieval(collection_name, search_query, top_k=10)
    return agent

# Agent 4: Contract Risk Analysis
def create_contract_risk_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    agent = Agent(
        model,
        system_prompt=(
            "You are a specialized contract risk analysis assistant. Your task is to analyze contract documents to identify clauses that could put the Company at a disadvantage. "
            "Follow these steps:\n"
            "1. Retrieve document sections that discuss contract terms, focusing on potential risks such as unilateral termination rights or other one-sided clauses using the provided tool if needed.\n"
            "2. Identify biased or risky clauses.\n"
            "3. Suggest modifications to balance the contract terms (for example, adding a notice period for termination).\n"
            "4. Present your output under the top-level heading **Analyzing Contract Risks** with clear markdown headings for both risks and recommendations.\n"
            "IMPORTANT: Do not include any function call syntax or tool formatting in your final response."
        )
    )
    @agent.tool
    async def search_document_content(context: RunContext[ContractRiskDeps], search_query: str) -> str:
        collection_name = context.deps.collection_name
        if not search_query or search_query.strip() == "":
            search_query = "contract risk analysis biased clauses unilateral termination rights modifications notice period"
        return await semantic_search_retrieval(collection_name, search_query, top_k=10)
    return agent

# Agent 5: Compliance Check
def create_compliance_agent():
    """Create and configure the compliance check agent"""
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    
    agent = Agent(
        model,
        system_prompt=(
            "You are a specialized compliance check assistant. Your task is to analyze RFP documents and determine "
            "whether the company meets all legal and domain eligibility requirements. Follow these steps:\n"
            "1. Analyze the document's requirements for eligibility: state registrations, certifications, experience.\n"
            "2. Compare the requirements against the company data provided.\n"
            "3. Calculate compliance scores for different categories and provide an overall weighted score.\n"
            "4. Identify any potential deal-breakers that would make the company ineligible to bid.\n"
            "5. Present your analysis in a clear, structured format with:\n"
            "   - A weighted scoring matrix showing compliance in different categories\n"
            "   - Identification of deal-breakers\n"
            "   - Recommendations for addressing compliance gaps\n\n"
            "Focus on accuracy and objectivity. Flag requirements that are unclear or not mentioned in the RFP."
        )
    )
    
    @agent.tool
    async def search_document_content(context: RunContext[ComplianceDeps], search_query: str) -> str:
        """
        Search the document for specific content related to compliance requirements.
        
        Args:
            context: The context with collection name.
            search_query: The specific query to search for in the document.
            
        Returns:
            The relevant document sections matching the search query.
        """
        collection_name = context.deps.collection_name
        
        if not search_query or search_query.strip() == "":
            # Default query if none provided
            search_query = "eligibility requirements compliance certifications registration state license experience"
        
        results = await semantic_search_retrieval(collection_name, search_query, top_k=10)
        return results
    
    @agent.tool
    async def get_company_data(context: RunContext[ComplianceDeps]) -> str:
        """
        Retrieve the company data for compliance checking.
        
        Args:
            context: The context with company data.
            
        Returns:
            Formatted string with company information.
        """
        company_data = context.deps.company_data
        
        if not company_data:
            return "Error: No company data available."
        
        result = "## Company Data for Compliance Check\n\n"
        
        # Format the company data for readability
        for key, value in company_data.items():
            result += f"**{key}**: {value}\n\n"
        
        return result
    
    return agent

# Run functions for each agent
async def run_summary_agent(collection_name: str) -> str:
    print(f"Generating summary for document collection: {collection_name}")
    agent = create_summary_agent()
    deps = SummaryDeps(collection_name=collection_name)
    summarization_prompt = (
        "Please provide a comprehensive summary of the document. First, retrieve the document content using the get_document_content tool, "
        "then analyze it to extract the overall context, scope of work, key points, and document type. "
        "Structure your response under the heading **Document Summary** with clear markdown formatting."
    )
    try:
        response = await agent.run(summarization_prompt, deps=deps)
        print("\n==== Document Summary ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running summary agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error generating summary: {str(e)}"

async def run_eligibility_agent(collection_name: str) -> str:
    print(f"\nExtracting eligibility criteria for document collection: {collection_name}")
    agent = create_eligibility_agent()
    deps = EligibilityDeps(collection_name=collection_name)
    eligibility_prompt = (
        "Extract all mandatory eligibility criteria from the document. First, retrieve the relevant sections using the search tool, "
        "then analyze them to summarize must-have qualifications, certifications, and experience needed to bid. "
        "Flag any missing requirements to prevent wasted effort on non-eligible proposals. "
        "Present your output under the heading **Extracting Mandatory Eligibility Criteria** using structured markdown."
    )
    try:
        response = await agent.run(eligibility_prompt, deps=deps)
        print("\n==== Extracting Mandatory Eligibility Criteria ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running eligibility agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error extracting eligibility criteria: {str(e)}"

async def run_submission_checklist_agent(collection_name: str) -> str:
    print(f"\nGenerating submission checklist for document collection: {collection_name}")
    agent = create_submission_checklist_agent()
    deps = SubmissionChecklistDeps(collection_name=collection_name)
    submission_prompt = (
        "Generate a submission checklist based on the document content. First, retrieve the sections related to document format and required attachments/forms using the search tool, "
        "then organize the information into a clear, structured checklist. "
        "Present your output under the heading **Generating a Submission Checklist** with clear markdown formatting."
    )
    try:
        response = await agent.run(submission_prompt, deps=deps)
        print("\n==== Generating a Submission Checklist ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running submission checklist agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error generating submission checklist: {str(e)}"

async def run_contract_risk_agent(collection_name: str) -> str:
    print(f"\nAnalyzing contract risks for document collection: {collection_name}")
    agent = create_contract_risk_agent()
    deps = ContractRiskDeps(collection_name=collection_name)
    contract_risk_prompt = (
        "Analyze the contract document to identify clauses that could put the Company at a disadvantage. First, retrieve the sections discussing contract terms using the search tool, "
        "then identify any biased or risky clauses and suggest modifications (for example, adding a notice period for termination). "
        "Present your output under the heading **Analyzing Contract Risks** with clear markdown headings for risks and recommendations."
    )
    try:
        response = await agent.run(contract_risk_prompt, deps=deps)
        print("\n==== Analyzing Contract Risks ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running contract risk agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error analyzing contract risks: {str(e)}"
    
async def run_compliance_agent(collection_name: str, company_data_file: str) -> str:
    """
    Run the compliance check agent on the specified collection.
    
    Args:
        collection_name: The collection to analyze.
        company_data_file: Path to the company data file.
        
    Returns:
        The compliance check results.
    """
    print(f"\nRunning compliance check for document collection: {collection_name}")
    
    # Load company data
    global COMPANY_DATA
    COMPANY_DATA = load_company_data(company_data_file)
    
    if not COMPANY_DATA:
        return "Error: Failed to load company data."
    
    print(f"Loaded company data: {len(COMPANY_DATA)} attributes found")
    
    agent = create_compliance_agent()
    deps = ComplianceDeps(collection_name=collection_name, company_data=COMPANY_DATA)
    
    # Multi-step prompt for comprehensive compliance check
    compliance_prompt = (
        "Perform a comprehensive compliance check to determine if the company is legally eligible to bid on this RFP. "
        "Follow these steps:\n\n"
        
        "1. First, use the search_document_content tool to search for 'eligibility requirements compliance certifications registration' "
        "to find the most relevant sections about legal requirements.\n\n"
        
        "2. Then search specifically for 'experience requirements domain technical capabilities' to understand "
        "domain-specific requirements.\n\n"
        
        "3. Use the get_company_data tool to retrieve information about the company's certifications, registrations, "
        "experience, and domain expertise.\n\n"
        
        "4. Create a weighted compliance scoring matrix with the following categories and weights:\n"
        "   - Legal Eligibility (40%): State registrations, basic qualifications\n"
        "   - Domain Match (25%): Alignment between company specialization and RFP requirements\n"
        "   - Certifications (20%): Required certifications and licenses\n"
        "   - Experience (15%): Years in business and past performance\n\n"
        
        "5. For each category, assign a score (0-100) based on how well the company meets the requirements, and calculate "
        "a weighted overall score.\n\n"
        
        "6. Identify any deal-breakers that would make the company ineligible to bid.\n\n"
        
        "7. Provide recommendations for addressing compliance gaps.\n\n"
        
        "Present your analysis in a clear, structured format with tables and sections. If the RFP doesn't mention specific "
        "requirements in some areas, note this in your analysis."
    )
    
    try:
        response = await agent.run(compliance_prompt, deps=deps)
        
        print("\n==== Compliance Check Results ====")
        print(response.content)
        
        return response.content
    except Exception as e:
        import traceback
        print(f"Error running compliance agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error performing compliance check: {str(e)}"

async def list_available_collections() -> List[str]:
    if not os.path.exists(FAISS_INDEX_FOLDER):
        return []
    return [d for d in os.listdir(FAISS_INDEX_FOLDER) if os.path.isdir(os.path.join(FAISS_INDEX_FOLDER, d))]

async def main_async():
    collections = await list_available_collections()
    if not collections:
        print("No document collections available. Please run rag_pipeline.py first to process documents.")
        return
    
    print("Available document collections:")
    for i, collection in enumerate(collections):
        print(f"{i+1}. {collection}")
    
    collection_idx = input("\nSelect collection number: ")
    try:
        idx = int(collection_idx) - 1
        if 0 <= idx < len(collections):
            collection_name = collections[idx]
        else:
            print("Invalid selection. Please enter a valid number.")
            return
    except ValueError:
        print("Invalid input. Please enter a number.")
        return
    
    print(f"\nSelected collection: {collection_name}")
    
    print("\nSelect an action:")
    print("1. Generate Document Summary")
    print("2. Extract Eligibility Criteria")
    print("3. Generate Submission Checklist")
    print("4. Analyze Contract Risks")
    print("5. Run Compliance Check")
    
    action = input("Enter your choice (1-5): ")
    
    if action == "1":
        await run_summary_agent(collection_name)
    elif action == "2":
        await run_eligibility_agent(collection_name)
    elif action == "3":
        await run_submission_checklist_agent(collection_name)
    elif action == "4":
        await run_contract_risk_agent(collection_name)
    elif action == "5":
        # Ask for company data file path
        company_data_file = input("\nEnter the path to company data file (PDF, DOCX, TXT, or JSON): ")
        if os.path.exists(company_data_file):
            await run_compliance_agent(collection_name, company_data_file)
        else:
            print(f"Error: The file '{company_data_file}' does not exist.")
    else:
        print("Invalid action. Please enter a number between 1 and 5.")

def main():
    asyncio.run(main_async())

if __name__ == "__main__":
    main()