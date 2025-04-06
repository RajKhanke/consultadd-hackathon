import pandas as pd
import PyPDF2
import os
from typing import Dict, List, Tuple, Optional
import json
import os
import sys
import asyncio
from dataclasses import dataclass
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
from pydantic import BaseModel, Field
import docx

from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from langchain.retrievers.document_compressors import CohereRerank
from langchain.retrievers import ContextualCompressionRetriever
from langchain.schema import Document

from pydantic_ai import Agent, RunContext
from pydantic_ai.models.groq import GroqModel
from pydantic_ai.providers.groq import GroqProvider

# Load environment variables
load_dotenv()

# Constants
FAISS_INDEX_FOLDER = os.path.join(os.path.expanduser("~"), "OneDrive", "Documents", "consultadd", "rag_data", "faiss_indexes")
MAX_CHUNKS_SUMMARY = 25  # Maximum chunks for summarization
MAX_CHUNKS_ELIGIBILITY = 50  # Maximum chunks for eligibility checks
RERANKING_ENABLED = False  # Set to True if you have a Cohere API key

# Check for required environment variables
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY not found in environment variables")

# Check for optional Cohere API key for reranking
COHERE_API_KEY = os.environ.get("COHERE_API_KEY")
if COHERE_API_KEY:
    RERANKING_ENABLED = True
    os.environ["COHERE_API_KEY"] = COHERE_API_KEY
else:
    print("Warning: COHERE_API_KEY not found. Reranking will be disabled.")

# Pydantic models for structured outputs
class DocumentSummary(BaseModel):
    """Document summary model"""
    context: str = Field(..., description="Overall context and purpose of the document")
    scope_of_work: str = Field(..., description="Description of the work or services being requested")
    key_points: List[str] = Field(..., description="Key points about the document")
    document_type: str = Field(..., description="Type of the document (RFP, proposal, contract, etc.)")

class EligibilityCriteria(BaseModel):
    """Eligibility criteria model"""
    mandatory_requirements: List[str] = Field(..., description="List of all mandatory requirements")
    qualifications: List[str] = Field(..., description="Required qualifications")
    certifications: List[str] = Field(..., description="Required certifications")
    experience: List[str] = Field(..., description="Required experience")
    missing_criteria: List[str] = Field(..., description="Potentially missing or unclear criteria")
    recommendation: str = Field(..., description="Overall assessment of the eligibility requirements")

# New Pydantic models for compliance check
class ComplianceCheckResult(BaseModel):
    """Compliance check result model"""
    overall_score: float = Field(..., description="Overall compliance score (0-100)")
    eligibility_score: float = Field(..., description="Legal eligibility score (0-100)")
    domain_match_score: float = Field(..., description="Domain alignment score (0-100)")
    certifications_score: float = Field(..., description="Certifications compliance score (0-100)")
    experience_score: float = Field(..., description="Experience requirements score (0-100)")
    deal_breakers: List[str] = Field(..., description="Potential deal-breakers identified")
    recommendations: List[str] = Field(..., description="Recommendations for compliance")

@dataclass
class ComplianceDeps:
    """Dependencies for the compliance check agent"""
    collection_name: str
    company_data: Dict[str, Any] 

COMPANY_DATA = {}

# Function to load and process company data
def load_company_data(file_path: str) -> Dict[str, Any]:
    """
    Load and process company data from the provided file.
    
    Args:
        file_path: Path to the company data file.
        
    Returns:
        Dictionary with company data.
    """
    # Determine file type and extract text
    ext = file_path.lower().rsplit(".", 1)[-1]
    raw_text = ""
    
    try:
        if ext == "pdf":
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        raw_text += page_text + "\n"
        elif ext == "docx":
            import docx
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    raw_text += paragraph.text + "\n"
                    
            # Handle tables in the document
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if len(row_text) >= 2:  # Ensure we have at least a key-value pair
                        raw_text += f"{row_text[0]}\n{row_text[1]}\n"
        
        elif ext == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()
        elif ext == "json":
            with open(file_path, "r", encoding="utf-8") as f:
                return json.load(f)
        else:
            raise ValueError(f"Unsupported file format: .{ext}")
    except Exception as e:
        print(f"Error loading company data: {str(e)}")
        return {}
    
    # Parse the text into key-value pairs
    company_data = {}
    
    # Split by lines and look for patterns that suggest field-value pairs
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    current_key = None
    
    for i, line in enumerate(lines):
        # Skip header-like lines
        if line.upper() == line and "DATA" in line.upper():
            continue
            
        if i < len(lines) - 1:
            # If this line looks like a field name and the next isn't a field name
            if not line.startswith("(") and not line.endswith(")") and not line.endswith(":"):
                current_key = line
                if i + 1 < len(lines):
                    company_data[current_key] = lines[i + 1]
    
    # If the above extraction didn't work well, try a simple key-value approach
    if len(company_data) < 5:  # Arbitrary threshold for successful extraction
        company_data = {}
        for i in range(0, len(lines)-1, 2):
            if i+1 < len(lines):
                if lines[i] and lines[i+1]:  # Both lines have content
                    company_data[lines[i]] = lines[i+1]
    
    # Add specific fields we know should be present based on the data provided
    expected_fields = [
        "Company Legal Name", "Principal Business Address", "State Registration Number",
        "Company Length of Existence", "Years of Experience in Temporary Staffing",
        "NAICS Codes", "State of Incorporation", "Services Provided",
        "Business Structure", "Licenses", "Historically Underutilized Business/DBE Status"
    ]
    
    # Ensure expected fields exist
    for field in expected_fields:
        if field not in company_data:
            # Try to find a close match
            for key in company_data.keys():
                if field.lower() in key.lower():
                    company_data[field] = company_data[key]
                    break
            
            # If still not found, mark as unknown
            if field not in company_data:
                company_data[field] = "Not provided in document"
    
    # Clean up any remaining issues
    final_data = {}
    for key, value in company_data.items():
        # Remove any extra whitespace or newlines
        cleaned_key = key.strip()
        cleaned_value = value.strip() if isinstance(value, str) else value
        
        # Skip empty keys/values
        if cleaned_key and cleaned_value:
            final_data[cleaned_key] = cleaned_value
    
    return final_data

# Document retrieval functions
async def retrieve_document_chunks(collection_name: str, is_summary: bool = True) -> str:
    """
    Retrieve document content based on task type.
    For summary: Uses first N chunks in order.
    """
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )
    index_path = os.path.join(FAISS_INDEX_FOLDER, collection_name)
    if not os.path.exists(index_path):
        return f"Error: FAISS index not found for collection '{collection_name}'."
    try:
        vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        docs = []
        for doc_id in vectorstore.docstore._dict:
            doc = vectorstore.docstore._dict[doc_id]
            if hasattr(doc, "page_content") and hasattr(doc, "metadata"):
                docs.append(doc)
        if is_summary:
            docs.sort(key=lambda x: x.metadata.get("chunk_id", 0))
            selected_docs = docs[:MAX_CHUNKS_SUMMARY]
        else:
            selected_docs = docs

        results = []
        for i, doc in enumerate(selected_docs):
            chunk_id = doc.metadata.get("chunk_id", i)
            source = doc.metadata.get("source", collection_name)
            results.append(f"## Document Section {i+1}\n**Source:** {source} | **Chunk:** {chunk_id}\n\n{doc.page_content}\n")
        return "\n\n".join(results)
    except Exception as e:
        import traceback
        print(f"Error retrieving document chunks: {str(e)}")
        print(traceback.format_exc())
        return f"Error retrieving document chunks: {str(e)}"

async def semantic_search_retrieval(collection_name: str, query: str, top_k: int = 10) -> str:
    """
    Perform semantic search on document collection and retrieve most relevant chunks.
    """
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )
    index_path = os.path.join(FAISS_INDEX_FOLDER, collection_name)
    if not os.path.exists(index_path):
        return f"Error: FAISS index not found for collection '{collection_name}'."
    try:
        vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        retriever_vectordb = vectorstore.as_retriever(search_kwargs={"k": top_k})
        all_docs = []
        for doc_id in vectorstore.docstore._dict:
            doc = vectorstore.docstore._dict[doc_id]
            if hasattr(doc, "page_content") and hasattr(doc, "metadata"):
                all_docs.append(Document(page_content=doc.page_content, metadata=doc.metadata))
        keyword_retriever = BM25Retriever.from_documents(all_docs)
        keyword_retriever.k = top_k
        ensemble_retriever = EnsembleRetriever(
            retrievers=[retriever_vectordb, keyword_retriever],
            weights=[0.7, 0.3]
        )
        if RERANKING_ENABLED:
            compressor = CohereRerank()
            final_retriever = ContextualCompressionRetriever(
                base_compressor=compressor,
                base_retriever=ensemble_retriever
            )
        else:
            final_retriever = ensemble_retriever
        docs = final_retriever.get_relevant_documents(query)
        results = []
        for i, doc in enumerate(docs):
            chunk_id = doc.metadata.get("chunk_id", i)
            source = doc.metadata.get("source", collection_name)
            results.append(f"## Document Section {i+1}\n**Source:** {source} | **Chunk:** {chunk_id}\n\n{doc.page_content}\n")
        if not results:
            return "No relevant document chunks found."
        return "\n\n".join(results)
    except Exception as e:
        import traceback
        print(f"Error in semantic search: {str(e)}")
        print(traceback.format_exc())
        return f"Error performing semantic search: {str(e)}"

# Dependency models for each agent
@dataclass
class SummaryDeps:
    collection_name: str

@dataclass
class EligibilityDeps:
    collection_name: str

@dataclass
class SubmissionChecklistDeps:
    collection_name: str

@dataclass
class ContractRiskDeps:
    collection_name: str

# Agent 1: Document Summary
def create_summary_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    # Added instruction to avoid any tool call formatting in the final answer.
    agent = Agent(
        model,
        system_prompt=(
              "You are a specialized document summarization assistant. Your task is to generate a comprehensive yet concise "
            "summary of a document. Follow these steps:\n"
            "1. Analyze the document content provided to understand its purpose and structure.\n"
            "2. Extract the core context, scope of work, and key points.\n"
            "3. Create a structured summary that includes:\n"
            "   - Overall context and purpose of the document\n"
            "   - Description of the work or services being requested\n"
            "   - Key points about requirements, deliverables, and important details\n"
            "   - Type of the document (RFP, contract, proposal, etc.)\n"
            "4. Present your summary in a clear, structured format.\n\n"
            "Your summary should be factual, objective, and based solely on the provided document content."
        )
    )
    @agent.tool
    async def get_document_content(context: RunContext[SummaryDeps]) -> str:
        collection_name = context.deps.collection_name
        return await retrieve_document_chunks(collection_name, is_summary=True)
    return agent

# Agent 2: Eligibility Criteria Extraction
def create_eligibility_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    agent = Agent(
        model,
        system_prompt=(
              "You are a specialized eligibility criteria extraction assistant. Your task is to analyze documents and identify "
            "all mandatory eligibility requirements. Follow these steps:\n"
            "1. Carefully review the document content provided.\n"
            "2. Extract and list all mandatory requirements, focusing on:\n"
            "   - Required qualifications and credentials\n"
            "   - Required certifications or licenses\n"
            "   - Minimum experience requirements\n"
            "   - Technical capabilities or resources needed\n"
            "3. Flag any missing or unclear eligibility criteria that should be specified but aren't.\n"
            "4. Provide a structured output with all mandatory requirements and flagged missing elements.\n"
            "5. Include a brief assessment of the overall eligibility requirements (strict, standard, minimal, etc.).\n\n"
            "Focus specifically on requirements using language like 'must have', 'required', 'shall', or 'mandatory'."
        )
    )
    @agent.tool
    async def search_document_content(context: RunContext[EligibilityDeps], search_query: str) -> str:
        collection_name = context.deps.collection_name
        if not search_query or search_query.strip() == "":
            search_query = "mandatory requirements qualifications certifications experience technical capabilities eligibility"
        return await semantic_search_retrieval(collection_name, search_query, top_k=10)
    return agent

# Agent 3: Submission Checklist Generation
def create_submission_checklist_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    agent = Agent(
        model,
        system_prompt=(
            "You are a specialized assistant for generating a submission checklist. Your task is to extract and structure the RFP submission requirements from a document. "
            "Follow these steps:\n"
            "1. Retrieve document sections relevant to submission requirements, including document format details (e.g., page limit, font type/size, line spacing, TOC requirements) and any specific attachments or forms using the provided tool if needed.\n"
            "2. Organize the extracted information into a structured checklist.\n"
            "3. Present your output under the top-level heading **Generating a Submission Checklist** with clear markdown headings and bullet points.\n"
            "IMPORTANT: Do not include any function call syntax or tool formatting in your final response."
        )
    )
    @agent.tool
    async def search_document_content(context: RunContext[SubmissionChecklistDeps], search_query: str) -> str:
        collection_name = context.deps.collection_name
        if not search_query or search_query.strip() == "":
            search_query = "submission checklist document format page limit font type line spacing attachments forms TOC requirements"
        return await semantic_search_retrieval(collection_name, search_query, top_k=10)
    return agent

# Agent 4: Contract Risk Analysis
def create_contract_risk_agent():
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY)
    )
    agent = Agent(
        model,
        system_prompt=(
            "You are a specialized contract risk analysis assistant. Your task is to analyze contract documents to identify clauses that could put the Company at a disadvantage. "
            "Follow these steps:\n"
            "1. Retrieve document sections that discuss contract terms, focusing on potential risks such as unilateral termination rights or other one-sided clauses using the provided tool if needed.\n"
            "2. Identify biased or risky clauses.\n"
            "3. Suggest modifications to balance the contract terms (for example, adding a notice period for termination).\n"
            "4. Present your output under the top-level heading **Analyzing Contract Risks** with clear markdown headings for both risks and recommendations.\n"
            "IMPORTANT: Do not include any function call syntax or tool formatting in your final response."
        )
    )
    @agent.tool
    async def search_document_content(context: RunContext[ContractRiskDeps], search_query: str) -> str:
        collection_name = context.deps.collection_name
        if not search_query or search_query.strip() == "":
            search_query = "contract risk analysis biased clauses unilateral termination rights modifications notice period"
        return await semantic_search_retrieval(collection_name, search_query, top_k=10)
    return agent

def create_compliance_agent():
    """Create and configure the compliance check agent with citations from the retrieved documents"""
    model = GroqModel(
        'llama-3.1-8b-instant',
        provider=GroqProvider(api_key=GROQ_API_KEY) # Set temperature to 0 to reduce hallucinations
    )
    
    agent = Agent(
        model,
        system_prompt=(
            "You are a specialized compliance check assistant tasked with determining if FirstStaff Workforce Solutions meets "
            "the eligibility requirements for a specific RFP. Your analysis must be based SOLELY on the company data and RFP requirements. "
            "For each assessment, include direct citations from the RFP requirements to justify your scoring. "
            "Provide a structured output with evidence-based scoring and rationale for each category."
        )
    )
    
    @agent.tool
    async def get_company_data(context: RunContext[ComplianceDeps]) -> str:
        """Retrieve the company data for compliance checking."""
        company_data = context.deps.company_data
        if not company_data:
            return '{"error": "No company data available."}'
        return json.dumps(company_data)
    
    @agent.tool
    async def get_rfp_requirements(context: RunContext[ComplianceDeps]) -> str:
        """Get all RFP requirements at once for compliance analysis."""
        collection_name = context.deps.collection_name
        search_query = (
            "mandatory requirements eligibility criteria certifications licenses experience "
            "state registration business structure NAICS codes qualifications"
        )
        # Use fewer chunks to prevent overwhelming the model
        results = await semantic_search_retrieval(collection_name, search_query, top_k=12)
        return results
    
    return agent

# Run functions for each agent
async def run_summary_agent(collection_name: str) -> str:
    print(f"Generating summary for document collection: {collection_name}")
    agent = create_summary_agent()
    deps = SummaryDeps(collection_name=collection_name)
    summarization_prompt = (
        "Please provide a comprehensive summary of the document. First, retrieve the document content using the get_document_content tool, "
        "then analyze it to extract the overall context, scope of work, key points, and document type. "
        "Structure your response under the heading **Document Summary** with clear markdown formatting."
    )
    try:
        response = await agent.run(summarization_prompt, deps=deps)
        print("\n==== Document Summary ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running summary agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error generating summary: {str(e)}"

async def run_eligibility_agent(collection_name: str) -> str:
    print(f"\nExtracting eligibility criteria for document collection: {collection_name}")
    agent = create_eligibility_agent()
    deps = EligibilityDeps(collection_name=collection_name)
    eligibility_prompt = (
        "Extract all mandatory eligibility criteria from the document. First, retrieve the relevant sections using the search tool, "
        "then analyze them to summarize must-have qualifications, certifications, and experience needed to bid. "
        "Flag any missing requirements to prevent wasted effort on non-eligible proposals. "
        "Present your output under the heading **Extracting Mandatory Eligibility Criteria** using structured markdown."
    )
    try:
        response = await agent.run(eligibility_prompt, deps=deps)
        print("\n==== Extracting Mandatory Eligibility Criteria ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running eligibility agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error extracting eligibility criteria: {str(e)}"

async def run_submission_checklist_agent(collection_name: str) -> str:
    print(f"\nGenerating submission checklist for document collection: {collection_name}")
    agent = create_submission_checklist_agent()
    deps = SubmissionChecklistDeps(collection_name=collection_name)
    submission_prompt = (
        "Generate a submission checklist based on the document content. First, retrieve the sections related to document format and required attachments/forms using the search tool, "
        "then organize the information into a clear, structured checklist. "
        "Present your output under the heading **Generating a Submission Checklist** with clear markdown formatting."
    )
    try:
        response = await agent.run(submission_prompt, deps=deps)
        print("\n==== Generating a Submission Checklist ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running submission checklist agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error generating submission checklist: {str(e)}"

async def run_contract_risk_agent(collection_name: str) -> str:
    print(f"\nAnalyzing contract risks for document collection: {collection_name}")
    agent = create_contract_risk_agent()
    deps = ContractRiskDeps(collection_name=collection_name)
    contract_risk_prompt = (
        "Analyze the contract document to identify clauses that could put the Company at a disadvantage. First, retrieve the sections discussing contract terms using the search tool, "
        "then identify any biased or risky clauses and suggest modifications (for example, adding a notice period for termination). "
        "Present your output under the heading **Analyzing Contract Risks** with clear markdown headings for risks and recommendations."
    )
    try:
        response = await agent.run(contract_risk_prompt, deps=deps, model_settings={"temperature": 0})
        print("\n==== Analyzing Contract Risks ====")
        print(response.data)
        return response.data
    except Exception as e:
        import traceback
        print(f"Error running contract risk agent: {str(e)}")
        print(traceback.format_exc())
        return f"Error analyzing contract risks: {str(e)}"
    
async def run_compliance_agent(collection_name: str, company_data_file: str) -> str:
    """Run the compliance check agent with citation support."""
    print(f"\nRunning comprehensive compliance check for document collection: {collection_name}")
    
    # Load company data
    global COMPANY_DATA
    COMPANY_DATA = load_company_data(company_data_file)
    
    if not COMPANY_DATA:
        return "Error: Failed to load company data."
    
    print(f"Loaded company data: {len(COMPANY_DATA)} attributes found")
    
    agent = create_compliance_agent()
    deps = ComplianceDeps(collection_name=collection_name, company_data=COMPANY_DATA)
    
    # Enhanced prompt to require specific citations
    prompt = (
        "Perform a thorough compliance check for FirstStaff Workforce Solutions against the RFP requirements. "
        "Follow these steps precisely:\n\n"
        
        "1. Use get_company_data to retrieve company information.\n"
        "2. Use get_rfp_requirements to retrieve all RFP requirements.\n"
        "3. For each of these categories, provide a detailed assessment with direct citations from the RFP:\n"
        "   - Legal Eligibility (40%): State registrations, business structure\n" 
        "   - Domain Match (25%): Alignment with RFP project scope\n"
        "   - Certifications (20%): Required certifications and licenses\n"
        "   - Experience (15%): Years of experience requirements\n"
        "4. For each category assessment, include:\n"
        "   - Specific RFP requirement citations (quoted text from the RFP)\n"
        "   - Corresponding company capabilities\n"
        "   - Justification for score assignment\n"
        "5. Identify any deal-breakers that disqualify the company.\n\n"
        
        "Return your analysis as valid JSON with this structure:\n"
        "{\n"
        "  \"overall_score\": float,\n"
        "  \"categories\": [\n"
        "    {\n"
        "      \"name\": string,\n"
        "      \"weight\": float,\n"
        "      \"score\": float,\n"
        "      \"rfp_requirements\": [list of requirement citations],\n"
        "      \"company_capabilities\": [list of relevant company data],\n"
        "      \"justification\": string\n"
        "    }\n"
        "  ],\n"
        "  \"deal_breakers\": [list of strings],\n"
        "  \"recommendations\": [list of strings],\n"
        "  \"eligibility_score\": float,\n"
        "  \"domain_match_score\": float,\n"
        "  \"certifications_score\": float,\n"
        "  \"experience_score\": float\n"
        "}\n\n"
        
        "IMPORTANT: Include direct citations from the RFP for every score assessment. Return only valid JSON."
    )
    
    try:
        print("Running compliance analysis with citations...")
        response = await agent.run(prompt, deps=deps, model_settings={"temperature": 0})
        result_data = response.data
        
        try:
            # Try to parse as JSON
            result_json = json.loads(result_data)
            
            # Ensure we have a simplified version that matches the previous output format
            simplified_json = {
                "overall_score": result_json.get("overall_score", 0),
                "eligibility_score": result_json.get("eligibility_score", 0),
                "domain_match_score": result_json.get("domain_match_score", 0),
                "certifications_score": result_json.get("certifications_score", 0),
                "experience_score": result_json.get("experience_score", 0),
                "deal_breakers": result_json.get("deal_breakers", []),
                "recommendations": result_json.get("recommendations", [])
            }
            
            # Print the detailed version with citations
            print("\n==== Detailed Compliance Check Results with Citations ====")
            print(json.dumps(result_json, indent=2))
            
            # Return the full detailed result
            return json.dumps(result_json, indent=2)
            
        except json.JSONDecodeError:
            # Extract JSON if there's surrounding text
            import re
            json_pattern = r'(\{[\s\S]*\})'
            matches = re.search(json_pattern, result_data)
            if matches:
                try:
                    json_str = matches.group(1)
                    result_json = json.loads(json_str)
                    print("\n==== Extracted Compliance Check Results ====")
                    print(json.dumps(result_json, indent=2))
                    return json.dumps(result_json, indent=2)
                except:
                    pass
            
            # If we can't extract JSON, return the raw result
            print("\n==== Compliance Check Results (Not JSON) ====")
            print(result_data)
            return result_data
            
    except Exception as e:
        import traceback
        print(f"Error running compliance agent: {str(e)}")
        print(traceback.format_exc())
        
        # Create a basic error response
        error_response = {
            "error": f"Compliance check failed: {str(e)}",
            "overall_score": 0,
            "eligibility_score": 0,
            "domain_match_score": 0, 
            "certifications_score": 0,
            "experience_score": 0,
            "deal_breakers": ["Compliance check failed to complete"],
            "recommendations": ["Try running eligibility check separately"]
        }
        
        return json.dumps(error_response)

async def list_available_collections() -> List[str]:
    if not os.path.exists(FAISS_INDEX_FOLDER):
        return []
    return [d for d in os.listdir(FAISS_INDEX_FOLDER) if os.path.isdir(os.path.join(FAISS_INDEX_FOLDER, d))]

async def main_async():
    collections = await list_available_collections()
    if not collections:
        print("No document collections available. Please run rag_pipeline.py first to process documents.")
        return
    
    print("Available document collections:")
    for i, collection in enumerate(collections):
        print(f"{i+1}. {collection}")
    
    collection_idx = input("\nSelect collection number: ")
    try:
        idx = int(collection_idx) - 1
        if 0 <= idx < len(collections):
            collection_name = collections[idx]
        else:
            print("Invalid selection. Please enter a valid number.")
            return
    except ValueError:
        print("Invalid input. Please enter a number.")
        return
    
    print(f"\nSelected collection: {collection_name}")
    
    print("\nSelect an action:")
    print("1. Generate Document Summary")
    print("2. Extract Eligibility Criteria")
    print("3. Generate Submission Checklist")
    print("4. Analyze Contract Risks")
    print("5. Run Compliance Check")
    
    action = input("Enter your choice (1-5): ")
    
    if action == "1":
        await run_summary_agent(collection_name)
    elif action == "2":
        await run_eligibility_agent(collection_name)
    elif action == "3":
        await run_submission_checklist_agent(collection_name)
    elif action == "4":
        await run_contract_risk_agent(collection_name)
    elif action == "5":
        # Ask for company data file path
        company_data_file = input("\nEnter the path to company data file (PDF, DOCX, TXT, or JSON): ")
        if os.path.exists(company_data_file):
            await run_compliance_agent(collection_name, company_data_file)
        else:
            print(f"Error: The file '{company_data_file}' does not exist.")
    else:
        print("Invalid action. Please enter a number between 1 and 5.")

def main():
    asyncio.run(main_async())

if __name__ == "__main__":
    main()